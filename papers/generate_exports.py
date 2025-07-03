#!/usr/bin/env python3
"""
Publication Export Generator
Generates multiple formats from BibTeX file: TXT, CSV, Markdown, DOCX
"""

import os
import csv
import json
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

def parse_bibtex_file(bib_file_path):
    """Parse BibTeX file and extract publication information"""
    publications = []
    
    with open(bib_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Find all BibTeX entries
    entries = re.findall(r'@(\w+)\{([^,]+),\s*(.*?)\n\}', content, re.DOTALL)
    
    for entry_type, entry_key, entry_content in entries:
        if entry_type.lower() in ['article', 'inproceedings', 'misc', 'phdthesis', 'techreport']:
            pub = {
                'type': entry_type,
                'key': entry_key,
                'title': '',
                'authors': '',
                'year': '',
                'venue': '',
                'doi': '',
                'url': '',
                'pages': '',
                'volume': '',
                'number': ''
            }
            
            # Extract fields using regex
            fields = re.findall(r'(\w+)\s*=\s*\{([^}]+)\}', entry_content)
            for field_name, field_value in fields:
                field_name = field_name.lower().strip()
                field_value = field_value.strip()
                
                if field_name == 'title':
                    pub['title'] = field_value
                elif field_name == 'author':
                    pub['authors'] = field_value
                elif field_name == 'year':
                    pub['year'] = field_value
                elif field_name in ['journal', 'booktitle', 'school']:
                    pub['venue'] = field_value
                elif field_name == 'doi':
                    pub['doi'] = field_value
                elif field_name == 'url':
                    pub['url'] = field_value
                elif field_name == 'pages':
                    pub['pages'] = field_value
                elif field_name == 'volume':
                    pub['volume'] = field_value
                elif field_name == 'number':
                    pub['number'] = field_value
            
            publications.append(pub)
    
    # Sort by year (descending) and then by title
    publications.sort(key=lambda x: (int(x['year']) if x['year'].isdigit() else 0, x['title']), reverse=True)
    
    return publications

def generate_text_format(publications, output_path):
    """Generate plain text format"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("SUSTAINABILITY LAB PUBLICATIONS\n")
        f.write("IIT Gandhinagar\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Total Publications: {len(publications)}\n\n")
        f.write("=" * 80 + "\n\n")
        
        current_year = None
        for i, pub in enumerate(publications, 1):
            if pub['year'] != current_year:
                current_year = pub['year']
                f.write(f"\n{current_year}\n")
                f.write("-" * len(current_year) + "\n\n")
            
            f.write(f"{i}. {pub['title']}\n")
            f.write(f"   Authors: {pub['authors']}\n")
            if pub['venue']:
                f.write(f"   Venue: {pub['venue']}\n")
            if pub['doi']:
                f.write(f"   DOI: {pub['doi']}\n")
            if pub['url']:
                f.write(f"   URL: {pub['url']}\n")
            f.write("\n")

def generate_csv_format(publications, output_path):
    """Generate CSV format"""
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Title', 'Authors', 'Year', 'Venue', 'Type', 'DOI', 'URL', 'Pages', 'Volume', 'Number'])
        
        for pub in publications:
            writer.writerow([
                pub['title'],
                pub['authors'],
                pub['year'],
                pub['venue'],
                pub['type'],
                pub['doi'],
                pub['url'],
                pub['pages'],
                pub['volume'],
                pub['number']
            ])

def generate_markdown_format(publications, output_path):
    """Generate Markdown format"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# Sustainability Lab Publications\n\n")
        f.write("**IIT Gandhinagar**  \n")
        f.write(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*  \n")
        f.write(f"*Total Publications: {len(publications)}*\n\n")
        f.write("---\n\n")
        
        current_year = None
        for pub in publications:
            if pub['year'] != current_year:
                current_year = pub['year']
                f.write(f"## {current_year}\n\n")
            
            f.write(f"### {pub['title']}\n\n")
            f.write(f"**Authors:** {pub['authors']}  \n")
            if pub['venue']:
                f.write(f"**Venue:** *{pub['venue']}*  \n")
            if pub['doi']:
                f.write(f"**DOI:** [{pub['doi']}](https://doi.org/{pub['doi']})  \n")
            if pub['url']:
                f.write(f"**URL:** [Link]({pub['url']})  \n")
            f.write("\n")

def generate_docx_format(publications, output_path):
    """Generate Word document format"""
    if not DOCX_AVAILABLE:
        print("Skipping DOCX generation - python-docx not available")
        return
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Sustainability Lab Publications', 0)
    title.alignment = 1  # Center alignment
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.add_run('IIT Gandhinagar').bold = True
    subtitle.alignment = 1
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
    meta.add_run(f'Total Publications: {len(publications)}')
    meta.alignment = 1
    
    doc.add_page_break()
    
    current_year = None
    for pub in publications:
        if pub['year'] != current_year:
            current_year = pub['year']
            doc.add_heading(current_year, level=1)
        
        # Title
        title_para = doc.add_paragraph()
        title_para.add_run(pub['title']).bold = True
        
        # Authors
        authors_para = doc.add_paragraph()
        authors_para.add_run('Authors: ').bold = True
        authors_para.add_run(pub['authors'])
        
        # Venue
        if pub['venue']:
            venue_para = doc.add_paragraph()
            venue_para.add_run('Venue: ').bold = True
            venue_para.add_run(pub['venue']).italic = True
        
        # DOI
        if pub['doi']:
            doi_para = doc.add_paragraph()
            doi_para.add_run('DOI: ').bold = True
            doi_para.add_run(pub['doi'])
        
        # URL
        if pub['url']:
            url_para = doc.add_paragraph()
            url_para.add_run('URL: ').bold = True
            url_para.add_run(pub['url'])
        
        doc.add_paragraph()  # Empty line
    
    doc.save(output_path)

def main():
    """Main function to generate all export formats"""
    script_dir = Path(__file__).parent
    bib_file = script_dir / 'bibtex' / 'all_publications.bib'
    exports_dir = script_dir / 'exports'
    
    # Create exports directory if it doesn't exist
    exports_dir.mkdir(exist_ok=True)
    
    if not bib_file.exists():
        print(f"Error: BibTeX file not found at {bib_file}")
        return
    
    print("Parsing BibTeX file...")
    publications = parse_bibtex_file(bib_file)
    print(f"Found {len(publications)} publications")
    
    # Generate all formats
    print("Generating text format...")
    generate_text_format(publications, exports_dir / 'all_publications.txt')
    
    print("Generating CSV format...")
    generate_csv_format(publications, exports_dir / 'all_publications.csv')
    
    print("Generating Markdown format...")
    generate_markdown_format(publications, exports_dir / 'all_publications.md')
    
    print("Generating Word document...")
    generate_docx_format(publications, exports_dir / 'all_publications.docx')
    
    print("\nExport generation complete!")
    print(f"Files saved to: {exports_dir}")
    print("Available formats:")
    print("  - all_publications.txt (Plain text)")
    print("  - all_publications.csv (Spreadsheet)")
    print("  - all_publications.md (Markdown)")
    if DOCX_AVAILABLE:
        print("  - all_publications.docx (Word document)")

if __name__ == "__main__":
    main()